# -*- coding: utf-8 -*-
"""
@Createtime: 2026-05-29
@description: GenerationContext - unified service injection for template handlers.

Instead of inheriting from BaseTemplateHandler, templates receive a
GenerationContext that provides all framework services. Templates implement
pure functions `preprocess(data, config)` and `generate(data, ctx)`.

Architecture:
    Template (pure functions)  ──ctx──▶  GenerationContext
                                               │
                    ┌──────────────────────────┼───────────────────────┐
                    ▼                          ▼                       ▼
            DocumentEditor          DocumentImageProcessor        DbDataReader
            TableProcessor          SummaryGenerator              TemplateInfo
"""

import os
import re
from copy import deepcopy
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .data_reader_db import DbDataReader
from .document_editor import DocumentEditor
from .document_image_processor import DocumentImageProcessor
from .handler_utils import TableProcessor
from .logger import setup_logger
from .summary_generator import SummaryGenerator, SummaryTemplates

_DEFAULT_LOGGER = setup_logger("GenerationContext")


# ═══════════════════════════════════════════════════════════════════
# Module-level utility functions (usable by preprocess without ctx)
# ═══════════════════════════════════════════════════════════════════

def set_default_dates(
    processed: Dict[str, Any],
    date_fields: List[str],
    date_format: str = "%Y-%m-%d",
) -> None:
    """Fill empty date fields with today's date.

    Args:
        processed: Data dict to mutate in-place.
        date_fields: List of field keys to check.
        date_format: strftime format string (default %Y-%m-%d).
    """
    today = datetime.now().strftime(date_format)
    for field in date_fields:
        if not processed.get(field) or processed.get(field) == 'today':
            processed[field] = today


def set_supplier_defaults(
    processed: Dict[str, Any],
    config: Dict[str, Any],
    supplier_fields: Optional[List[str]] = None,
) -> str:
    """Set supplier name default and optionally fill related fields.

    Args:
        processed: Data dict to mutate in-place.
        config: Global config dict (source of supplierName).
        supplier_fields: Extra fields to fill with the supplier name.

    Returns:
        The resolved supplier name.
    """
    supplier = processed.get('supplier_name') or config.get('supplierName', '')
    processed['supplier_name'] = supplier
    if supplier_fields:
        for field in supplier_fields:
            if not processed.get(field):
                processed[field] = supplier
    return supplier


def gen_report_id(
    prefix: str = "RPT",
    date_format: str = "%Y%m%d",
    random_length: int = 4,
    use_sequence: bool = False,
) -> str:
    """Generate a unique report ID.

    Args:
        prefix: ID prefix (e.g. 'RPT', 'YHBH', 'IR').
        date_format: strftime format for the date portion.
        random_length: Number of random alphanumeric chars for the suffix.
        use_sequence: If True, use time-based sequence instead of random chars.

    Returns:
        Formatted ID string, e.g. 'RPT-20260529-ABCD'.
    """
    import random as _random
    import string
    import time

    date_str = datetime.now().strftime(date_format)
    if use_sequence:
        suffix = f"{int(time.time()) % 10000:04d}"
    else:
        suffix = ''.join(
            _random.choices(string.ascii_uppercase + string.digits, k=random_length)
        )
    return f"{prefix}-{date_str}-{suffix}"


class GenerationContext:
    """
    Unified service context injected into template handler functions.

    Provides all framework services that templates need:
    - Document loading, editing, saving
    - Image processing
    - Table population
    - TOC insertion
    - Vulnerability lookup
    - Summary generation
    - Logging (TXT + SQLite)
    - Utility helpers (date, ID generation, path building)

    Usage in template handler.py:
        def generate(data: dict, ctx: GenerationContext) -> Tuple[bool, str, str]:
            doc = ctx.load_document()
            ctx.replace_text(replacements)
            ctx.process_single_image('#placeholder#', data.get('image'))
            return True, ctx.save(filename), "Report generated"
    """

    def __init__(
        self,
        template_dir: str,
        template_info: Any,
        config: Optional[Dict[str, Any]] = None,
        output_dir: str = "",
        logger_instance: Any = None,
    ):
        """
        Initialize GenerationContext.

        Args:
            template_dir: Absolute path to template directory
            template_info: TemplateInfo Pydantic model
            config: Global configuration dict
            output_dir: Base output directory for generated reports
            logger_instance: Logger instance (defaults to 'GenerationContext')
        """
        self.template_dir = template_dir
        self.template_info = template_info
        self.config = config or {}
        self.output_dir = output_dir
        self.logger = logger_instance or _DEFAULT_LOGGER

        # Lazy-initialized services
        self._doc: Optional[DocxDocument] = None
        self._editor: Optional[DocumentEditor] = None
        self._img_processor: Optional[DocumentImageProcessor] = None
        self._db_reader: Optional[DbDataReader] = None

        # Cached template path
        self._template_path: Optional[str] = None

    # ── Logger access ──────────────────────────────────────────────

    def setup_logger(self, name: str) -> Any:
        """Create a named logger for the template."""
        return setup_logger(name)

    # ── Document loading ───────────────────────────────────────────

    def load_document(self) -> Optional[DocxDocument]:
        """
        Load the template .docx file.

        Returns:
            Document object or None if template not found.
        """
        template_path = self._get_template_path()
        if template_path and os.path.exists(template_path):
            self._doc = Document(template_path)
            # Reset derived services since doc changed
            self._editor = None
            self._img_processor = None
            return self._doc
        return None

    @property
    def doc(self) -> Optional[DocxDocument]:
        """Access the currently loaded document (lazy-loads on first access)."""
        if self._doc is None:
            self.load_document()
        return self._doc

    @property
    def editor(self) -> DocumentEditor:
        """Access the DocumentEditor for the current document."""
        if self._editor is None:
            if self._doc is None:
                self.load_document()
            if self._doc is None:
                raise RuntimeError("Cannot create editor: document not loaded")
            self._editor = DocumentEditor(self._doc)
        return self._editor

    @property
    def img_processor(self) -> DocumentImageProcessor:
        """Access the DocumentImageProcessor for the current document."""
        if self._img_processor is None:
            if self._doc is None:
                self.load_document()
            if self._doc is None:
                raise RuntimeError("Cannot create image processor: document not loaded")
            self._img_processor = DocumentImageProcessor(self._doc, [])
        return self._img_processor

    # ── Text replacement ───────────────────────────────────────────

    def replace_text(
        self,
        replacements: Dict[str, str],
        enable_risk_color: bool = False,
        risk_key: Optional[str] = None,
    ) -> None:
        """
        Replace placeholders in the document with text values.

        Args:
            replacements: Dict mapping placeholder to value, e.g. {'#name#': 'Alice'}
            enable_risk_color: If True, apply risk-level colors to the risk placeholder
            risk_key: Risk level placeholder key (default "#overall_risk_level#")
        """
        kwargs = {'enable_risk_color': enable_risk_color}
        if risk_key is not None:
            kwargs['risk_key'] = risk_key
        self.editor.replace_report_text(replacements, **kwargs)

    def replace_text_colored(self, replacements: Dict[str, str]) -> None:
        """Replace text with risk-level colors enabled (convenience)."""
        self.replace_text(replacements, enable_risk_color=True)

    # ── Image processing ───────────────────────────────────────────

    def process_single_image(
        self,
        placeholder: str,
        image_data: Any,
        fallback_text: str = "（未提供）",
    ) -> None:
        """
        Process a single image placeholder.

        Args:
            placeholder: Placeholder text (e.g. '#screenshot#')
            image_data: Image path string or dict with 'path' key
            fallback_text: Text to show when no image is provided
        """
        if image_data:
            img_path = image_data if isinstance(image_data, str) else image_data.get('path', '')
            if img_path and os.path.exists(img_path):
                self.img_processor.replace_placeholder_with_images(
                    placeholder, [{'path': img_path, 'description': ''}]
                )
                return
        # No image: clear placeholder
        self.img_processor.replace_placeholder_with_images(placeholder, [])

    def process_image_list(
        self,
        placeholder: str,
        images: List[Any],
        keyword: Optional[str] = None,
    ) -> None:
        """
        Process a list of images for a placeholder.

        Args:
            placeholder: Placeholder text
            images: List of image paths (str) or dicts with 'path'/'description'
            keyword: Optional alternative keyword for placeholder location
        """
        target_keyword = keyword or placeholder
        doc = self.doc
        if doc is None:
            return

        cleanup_tokens = [placeholder]
        if keyword and keyword != placeholder:
            cleanup_tokens.append(keyword)

        if not images:
            for token in cleanup_tokens:
                self.img_processor.replace_placeholder_with_images(token, [])
            return

        # Normalize image data
        normalized: List[Dict[str, str]] = []
        for item in images:
            if isinstance(item, dict):
                img_path = item.get('path', '')
                description = item.get('description', '') or item.get('desc', '')
            else:
                img_path = str(item)
                description = ''
            if img_path:
                normalized.append({'path': img_path, 'description': description})

        if not normalized:
            for token in cleanup_tokens:
                self.img_processor.replace_placeholder_with_images(token, [])
            return

        # Try table cell first
        target_cell = None
        target_cell_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target_keyword in cell.text:
                        target_cell = cell
                        target_cell_text = cell.text
                        break
                if target_cell:
                    break
            if target_cell:
                break

        if target_cell:
            compact = ''.join(target_cell_text.split())
            is_placeholder_only = compact in {
                ''.join(placeholder.split()),
                ''.join(target_keyword.split()),
            }
            if is_placeholder_only:
                for para in target_cell.paragraphs:
                    if target_keyword in para.text:
                        para.text = para.text.replace(target_keyword, '')
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, '')
                self.img_processor.insert_images_into_cell(target_cell, normalized)
            else:
                self.img_processor.replace_placeholder_with_images(target_keyword, normalized)
            return

        # Fallback: paragraph-based placeholder
        target_para_text = None
        for para in doc.paragraphs:
            if target_keyword in para.text:
                target_para_text = para.text
                break

        if target_para_text is not None:
            self.img_processor.replace_placeholder_with_images(target_keyword, normalized)
            return

        # Cleanup if target not found
        for token in cleanup_tokens:
            self.img_processor.replace_placeholder_with_images(token, [])

    def insert_images_into_cell(
        self,
        cell: Any,
        image_list: List[Any],
        max_width_inches: float = 5.5,
    ) -> None:
        """Insert multiple images directly into a table cell."""
        # Normalize
        normalized: List[Dict[str, str]] = []
        for item in image_list:
            if isinstance(item, dict):
                img_path = item.get('path', '')
                description = item.get('description', '') or item.get('desc', '')
            else:
                img_path = str(item)
                description = ''
            if img_path:
                normalized.append({'path': img_path, 'description': description})
        self.img_processor.insert_images_into_cell(cell, normalized, max_width_inches)

    def insert_image_run(
        self,
        paragraph: Any,
        img_path: str,
        max_width_inches: float = 6.0,
    ) -> Any:
        """Insert a single image into a paragraph, returning the run."""
        return self.img_processor.insert_image_run(paragraph, img_path, max_width_inches)

    # ── Table operations ───────────────────────────────────────────

    def populate_table(
        self,
        table_header_text: str,
        data_rows: List[Dict[str, Any]],
        row_builder_func: Callable,
        keep_header_rows: int = 1,
        clear_indent: bool = False,
    ) -> bool:
        """
        Populate a table found by header text with data rows.

        Args:
            table_header_text: Text to find the table by (first cell match)
            data_rows: List of data dicts to populate
            row_builder_func: Function(row, data_item) -> None to build each row
            keep_header_rows: Number of header rows to preserve
            clear_indent: If True, clear paragraph indent in all cells

        Returns:
            True if table found and populated, False otherwise
        """
        doc = self.doc
        if doc is None:
            return False
        return TableProcessor.populate_table(
            doc,
            table_header_text,
            data_rows,
            row_builder_func,
            keep_header_rows=keep_header_rows,
            clear_indent=clear_indent,
            logger_instance=self.logger,
        )

    @staticmethod
    def clear_paragraph_indent(para: Any) -> None:
        """Clear first-line indent on a paragraph."""
        DocumentEditor.clear_paragraph_indent(para)

    # ── TOC ────────────────────────────────────────────────────────

    def insert_toc(
        self,
        placeholder: str = "#toc#",
        toc_title: str = "目  录",
    ) -> bool:
        """
        Insert a Table of Contents field at the placeholder location.

        Args:
            placeholder: The TOC placeholder text
            toc_title: Title text for the TOC section

        Returns:
            True if TOC was inserted, False if placeholder not found
        """
        return self.editor.insert_toc_at_placeholder(placeholder, toc_title)

    # ── Save / output ──────────────────────────────────────────────

    def save(self, filename: str) -> str:
        """
        Save the document to the output directory with the given filename.

        This is the primary save method for templates - it combines path
        building and document saving in one call.

        Args:
            filename: Output filename (e.g. '【海淀】【漏洞】Report【高危】.docx')

        Returns:
            Absolute path to the saved file
        """
        output_path = os.path.join(self.output_dir, filename)
        return self.save_document(self.doc, output_path)

    def save_document(
        self,
        doc: DocxDocument,
        output_path: str,
    ) -> str:
        """
        Save a document, handling filename conflicts.

        Args:
            doc: Document object
            output_path: Target path

        Returns:
            Actual saved file path (may differ if conflict resolved)
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        if os.path.exists(output_path):
            base, ext = os.path.splitext(output_path)
            count = 1
            while os.path.exists(f"{base}-{count}{ext}"):
                count += 1
            output_path = f"{base}-{count}{ext}"

        doc.save(output_path)
        return output_path

    def build_output_path(
        self,
        unit_name: str,
        filename: str,
    ) -> str:
        """
        Build a safe output file path under a unit subdirectory.

        Args:
            unit_name: Organization/unit name (creates subdirectory)
            filename: Desired filename

        Returns:
            Full output path
        """
        safe_unit = self.sanitize_filename(str(unit_name).strip('. '))
        if not safe_unit:
            safe_unit = "Unknown"
        safe_filename = self.sanitize_filename(str(filename).strip())
        if not safe_filename:
            safe_filename = f"report_{self.get_date()}.docx"

        company_dir = os.path.join(self.output_dir, safe_unit)
        os.makedirs(company_dir, exist_ok=True)

        return os.path.join(company_dir, safe_filename)

    # ── Replacements building ──────────────────────────────────────

    def build_replacements(
        self,
        data: Dict[str, Any],
        extra: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, str]:
        """
        Build a replacements dict from template field definitions + extra data.

        Args:
            data: Form data dict
            extra: Additional key-value pairs to include

        Returns:
            Dict mapping placeholders to string values
        """
        replacements: Dict[str, str] = {}

        if self.template_info:
            for field_def in self.template_info.fields:
                key = field_def.template_placeholder or f"#{field_def.key}#"
                value = data.get(field_def.key)
                if value is None:
                    value = field_def.default if field_def.default != 'today' else ''

                # Skip image/complex types for text replacement
                if field_def.type in (
                    'image', 'image_list', 'grouped_image_list'
                ):
                    continue
                if isinstance(value, list):
                    continue

                replacements[key] = str(value) if value is not None else ""

        if extra:
            for k, v in extra.items():
                if not k.startswith("#"):
                    k = f"#{k}#"
                replacements[k] = str(v) if v is not None else ""

        return replacements

    # ── Utility helpers ────────────────────────────────────────────

    def get_date(self, format_str: str = "%Y-%m-%d") -> str:
        """Get current date formatted string."""
        return datetime.now().strftime(format_str)

    def gen_id(
        self,
        prefix: str = "RPT",
        date_format: str = "%Y%m%d",
        random_length: int = 4,
        use_sequence: bool = False,
    ) -> str:
        """
        Generate a report ID. Delegates to the module-level gen_report_id().

        Args:
            prefix: ID prefix
            date_format: Date format string
            random_length: Random suffix length
            use_sequence: If True, use time-based sequence instead of random

        Returns:
            Formatted ID string, e.g. 'RPT-20260529-ABCD'
        """
        return gen_report_id(prefix, date_format, random_length, use_sequence)

    def sanitize_filename(self, filename: str) -> str:
        """Remove illegal characters from a filename."""
        return re.sub(r'[<>:"/\\|?*]', '_', filename)

    def create_output_dir(self, base_dir: str, sub_dir: str = "") -> str:
        """Create and return a safe output directory."""
        os.makedirs(base_dir, exist_ok=True)
        if not sub_dir:
            return base_dir
        safe = self.sanitize_filename(str(sub_dir).strip('. '))
        if not safe:
            safe = "Unknown"
        out = os.path.join(base_dir, safe)
        os.makedirs(out, exist_ok=True)
        return out

    # ── Vulnerability lookup ───────────────────────────────────────

    def _get_db_reader(self) -> DbDataReader:
        """Lazy-init DbDataReader from config."""
        if self._db_reader is None:
            db_path_config = self.config.get("vul_or_icp", "data/combined.db")
            if not os.path.isabs(db_path_config):
                # template_dir = backend/templates/{id}/, go up 2 to backend/
                backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(self.template_dir)))
                db_path = os.path.join(backend_dir, db_path_config)
            else:
                db_path = db_path_config
            self._db_reader = DbDataReader(db_path, "", "")
        return self._db_reader

    def lookup_vulnerability(self, vuln_id: str) -> Dict[str, Any]:
        """
        Look up a vulnerability by ID or name from the database.

        Args:
            vuln_id: Vulnerability ID (hash) or name

        Returns:
            Vulnerability info dict, or empty dict if not found
        """
        try:
            db = self._get_db_reader()
            _, vulns = db.read_vulnerabilities_from_db()

            if vuln_id in vulns:
                return vulns[vuln_id]

            # Try case-insensitive name match
            for v in vulns.values():
                name = v.get('Vuln_Name', '')
                if name and name.lower() == vuln_id.lower():
                    return v

            return {}
        except Exception as e:
            self.logger.error(f"Vulnerability lookup failed for '{vuln_id}': {e}")
            return {}

    def get_vulnerability_name(self, vuln_id: str) -> str:
        """
        Get a vulnerability's display name by ID.

        Args:
            vuln_id: Vulnerability ID

        Returns:
            Display name or empty string
        """
        vuln = self.lookup_vulnerability(vuln_id)
        return vuln.get('Vuln_Name', '')

    # ── Summary generation ─────────────────────────────────────────

    @staticmethod
    def summarize_count(
        items: List[Dict[str, Any]],
        type_key: str,
        type_names: Dict[str, str],
        template_zero: str,
        template_single: str,
        template_multi: str,
        connector: str = '、',
        last_connector: str = '以及',
    ) -> str:
        """Generate a count-based summary description."""
        return SummaryGenerator.count_summary(
            items=items,
            type_key=type_key,
            type_names=type_names,
            template_zero=template_zero,
            template_single=template_single,
            template_multi=template_multi,
            connector=connector,
            last_connector=last_connector,
        )

    @staticmethod
    def summarize_data(
        items: List[Dict[str, Any]],
        type_key: str,
        count_key: str,
        template_zero: str,
        template_with_data: str,
        connector: str = '，',
    ) -> Tuple[str, int]:
        """Generate a data-quantity summary description."""
        return SummaryGenerator.data_summary(
            items=items,
            type_key=type_key,
            count_key=count_key,
            template_zero=template_zero,
            template_with_data=template_with_data,
            connector=connector,
        )

    @property
    def summary_templates(self) -> type:
        """Access to predefined SummaryTemplates configs."""
        return SummaryTemplates

    # ── Fallback report ────────────────────────────────────────────

    def generate_fallback(
        self,
        data: Dict[str, Any],
    ) -> Tuple[bool, str, str]:
        """
        Generate a fallback report when template.docx is missing.

        Creates a simple document from schema field definitions.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            # Title
            title_text = self.template_info.name if self.template_info else "Report"
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Fields by group
            if self.template_info and self.template_info.field_groups:
                for group in self.template_info.field_groups:
                    doc.add_heading(group.name, level=1)
                    fields = [
                        f for f in self.template_info.fields if f.group == group.id
                    ]
                    fields.sort(key=lambda x: x.order)
                    if fields:
                        table = doc.add_table(rows=len(fields), cols=2)
                        table.style = 'Table Grid'
                        for i, field in enumerate(fields):
                            row = table.rows[i]
                            row.cells[0].text = field.label
                            value = data.get(field.key, '')
                            if isinstance(value, list):
                                row.cells[1].text = f"[Contains {len(value)} items]"
                            else:
                                row.cells[1].text = str(value)
            else:
                doc.add_heading("Report Data", level=1)
                for k, v in data.items():
                    doc.add_paragraph(f"{k}: {v}")

            output_path = self._generate_output_path_from_template(data)
            final_path = self.save_document(doc, output_path)
            return True, final_path, "Fallback report generated (template missing)"
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, "", f"Fallback generation failed: {str(e)}"

    def _generate_output_path_from_template(self, data: Dict[str, Any]) -> str:
        """Generate output path using template output config."""
        if not self.template_info:
            return os.path.join(self.output_dir, "report.docx")

        output_config = self.template_info.output_config
        filename_pattern = output_config.get(
            'filename_pattern', '{vul_name}_{date}.docx'
        )
        output_dir_pattern = output_config.get('output_dir', '')

        now = datetime.now()
        reps = {
            'date': now.strftime('%Y-%m-%d'),
            'datetime': now.strftime('%Y%m%d_%H%M%S'),
            'timestamp': str(int(now.timestamp())),
        }
        reps.update(data)

        filename = filename_pattern
        for k, v in reps.items():
            filename = filename.replace(f'{{{k}}}', str(v) if v else '')
        filename = self.sanitize_filename(filename)

        if output_dir_pattern:
            out_dir = output_dir_pattern
            for k, v in reps.items():
                out_dir = out_dir.replace(f'{{{k}}}', str(v) if v else '')
            out_dir = os.path.join(self.output_dir, out_dir)
        else:
            out_dir = self.output_dir

        os.makedirs(out_dir, exist_ok=True)
        return os.path.join(out_dir, filename)

    # ── Logging (TXT + SQLite postprocess) ─────────────────────────

    def write_txt_log(
        self,
        log_prefix: str,
        fields: List[Any],
    ) -> None:
        """
        Write a TXT format log entry.

        Args:
            log_prefix: Log file prefix (e.g. 'vuln', 'penetration')
            fields: Field values to record
        """
        report_date = self.get_date()
        log_file = os.path.join(
            self.output_dir, f"{report_date}_{log_prefix}_output.txt"
        )
        log_line = "\t".join([str(f) if f else '' for f in fields])
        with open(log_file, 'a+', encoding='utf-8') as f:
            f.write('\n' + log_line)

    def write_db_log(
        self,
        db_name: str,
        table_name: str,
        record: Dict[str, Any],
    ) -> None:
        """
        Write a record to SQLite database log.

        Args:
            db_name: Database filename (e.g. 'report_log.db')
            table_name: Table name
            record: Dict of column -> value
        """
        import sqlite3

        db_path = os.path.join(self.output_dir, db_name)
        columns = list(record.keys())
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        try:
            cursor.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                (table_name,),
            )
            table_exists = cursor.fetchone()

            if not table_exists:
                cols_def = ', '.join([f'"{col}" TEXT' for col in columns])
                cursor.execute(f'CREATE TABLE {table_name} ({cols_def})')
            else:
                cursor.execute(f"PRAGMA table_info({table_name})")
                existing = [row[1] for row in cursor.fetchall()]
                for col in columns:
                    if col not in existing:
                        cursor.execute(
                            f'ALTER TABLE {table_name} ADD COLUMN "{col}" TEXT'
                        )

            cols_sql = ', '.join([f'"{col}"' for col in columns])
            placeholders = ', '.join(['?' for _ in columns])
            values = [str(v) if v is not None else '' for v in record.values()]
            cursor.execute(
                f'INSERT INTO {table_name} ({cols_sql}) VALUES ({placeholders})',
                values,
            )
            conn.commit()
        except Exception as e:
            self.logger.error(f"DB log error: {e}")
        finally:
            conn.close()

    def postprocess(
        self,
        output_path: str,
        data: Dict[str, Any],
        log_prefix: str = "",
        log_fields: Optional[List[str]] = None,
        db_table: str = "",
        db_name: str = "",
        db_field_map: Optional[Dict[str, str]] = None,
    ) -> None:
        """
        Run post-generation logging (TXT + SQLite).

        Args:
            output_path: Path to the generated report
            data: Report data
            log_prefix: TXT log file prefix
            log_fields: List of data field keys to log to TXT
            db_table: Database table name
            db_name: Database filename
            db_field_map: Mapping of db_column -> data_key for DB record
        """
        try:
            report_date = (
                data.get('report_date')
                or data.get('report_time')
                or self.get_date()
            )

            # TXT log
            if log_prefix and log_fields:
                field_values = [str(data.get(k, '')) for k in log_fields]
                self.write_txt_log(log_prefix, field_values)

            # DB log
            if db_table and db_name and db_field_map:
                record = {
                    db_col: data.get(data_key, '')
                    for db_col, data_key in db_field_map.items()
                }
                self.write_db_log(db_name, db_table, record)
        except Exception as e:
            self.logger.error(f"Postprocess error: {e}")

    # ── Internal helpers ───────────────────────────────────────────

    def _get_template_path(self) -> Optional[str]:
        """Resolve the template .docx file path."""
        if self._template_path is None:
            template_file = os.path.join(
                self.template_dir, self.template_info.template_file
            )
            if os.path.exists(template_file):
                self._template_path = template_file
        return self._template_path

    def _reload_document(self) -> Optional[DocxDocument]:
        """Force reload the document (clears cached state)."""
        self._doc = None
        self._editor = None
        self._img_processor = None
        return self.load_document()
