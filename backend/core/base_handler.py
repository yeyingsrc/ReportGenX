# -*- coding: utf-8 -*-
"""
@Createtime: 2026-01-24
@description: 模板处理器基类 - 定义报告生成的标准接口
所有模板处理器都应该继承此基类
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional, Tuple, TypedDict
from datetime import datetime
from docx import Document
from docx.document import Document as DocxDocument

from .handler_registry import HandlerRegistry
from .logger import setup_logger

# 初始化日志记录器
logger = setup_logger('BaseHandler')


# ========== 类型定义 ==========
class GenerateResult(TypedDict):
    """报告生成结果类型"""
    success: bool
    report_path: str
    message: str
    errors: List[str]


class ValidationResult(TypedDict):
    """验证结果类型"""
    valid: bool
    errors: List[str]


class BaseTemplateHandler(ABC):
    """
    模板处理器基类
    
    每种报告模板都需要一个对应的 Handler 类来处理其特定的业务逻辑。
    Handler 负责：
    1. 数据预处理 (preprocess)
    2. 数据验证 (validate)
    3. 报告生成 (generate)
    4. 后处理 (postprocess)
    """
    
    def __init__(self, template_dir: str, template_info: Any, config: Optional[Dict[str, Any]] = None):
        """
        初始化处理器
        
        Args:
            template_dir: 模板目录绝对路径 (e.g. backend/templates/vuln_report)
            template_info: TemplateInfo Pydantic 模型实例
            config: 全局配置
        """
        self.template_dir = template_dir
        self.template_id = template_info.id
        self.template_info = template_info
        self.config = config or {}
        self.template_manager = None  # 解耦后不再依赖 TemplateManager 实例
    
    @abstractmethod
    def preprocess(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        数据预处理
        
        在生成报告前对数据进行预处理，如：
        - 自动生成编号
        - 日期格式化
        - 数据补全
        - 计算派生字段
        
        Args:
            data: 原始表单数据
            
        Returns:
            预处理后的数据
        """
        pass
    
    def validate(self, data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """
        数据验证
        
        验证数据是否符合模板要求，基于 self.template_info 的字段定义。
        子类可以覆盖此方法添加自定义验证
        
        Args:
            data: 表单数据
            
        Returns:
            (是否有效, 错误信息列表)
        """
        if not self.template_info:
            return True, []
        
        errors = []
        
        # 检查必填字段
        for field_def in self.template_info.fields:
            if field_def.required:
                value = data.get(field_def.key, "")
                if not value or (isinstance(value, str) and not value.strip()):
                    errors.append(f"字段 '{field_def.label}' 为必填项")
            
            # 检查字段验证规则
            if field_def.validation:
                pattern = field_def.validation.get('pattern')
                if pattern:
                    value = data.get(field_def.key, "")
                    if value and not re.match(pattern, str(value)):
                        errors.append(field_def.validation.get('message', f"字段 '{field_def.label}' 格式不正确"))
        
        # 检查全局验证规则
        for rule in self.template_info.validation_rules:
            if rule.rule == 'required':
                for field_key in rule.fields:
                    value = data.get(field_key, "")
                    if not value or (isinstance(value, str) and not value.strip()):
                        errors.append(rule.message)
                        break
        
        return len(errors) == 0, errors
    
    @abstractmethod
    def generate(self, data: Dict[str, Any], output_dir: str) -> Tuple[bool, str, str]:
        """
        生成报告
        
        核心方法，负责：
        1. 加载 docx 模板
        2. 执行占位符替换
        3. 插入图片
        4. 保存文件
        
        Args:
            data: 预处理后的表单数据
            output_dir: 输出目录
            
        Returns:
            (成功标志, 输出文件路径, 消息)
        """
        pass
    
    def postprocess(self, output_path: str, data: Dict[str, Any]) -> None:
        """
        后处理 - 使用模板方法模式
        
        报告生成后的处理，包括：
        - 记录 TXT 日志
        - 写入 SQLite 数据库
        
        子类只需实现 _get_log_fields(), _get_log_prefix(), 
        _build_db_record(), _get_db_table_name() 即可
        
        Args:
            output_path: 生成的报告文件路径
            data: 表单数据
        """
        try:
            report_date = data.get('report_date') or data.get('report_time') or self.get_current_date()
            
            # 获取子类定义的日志字段和前缀
            log_fields = self._get_log_fields(data, report_date)
            log_prefix = self._get_log_prefix()
            
            # 写入 TXT 日志
            if log_fields and log_prefix:
                self.write_txt_log(self.output_dir, log_prefix, log_fields)
            
            # 写入 SQLite 数据库
            db_table = self._get_db_table_name()
            if db_table:
                db_name = f"{report_date}_output.db"
                record = self._build_db_record(data, report_date, output_path)
                if record:
                    self.write_db_log(self.output_dir, db_name, db_table, record)
                    
        except Exception as e:
            logger.error(f"Postprocess error: {e}")
    
    def _get_log_fields(self, data: Dict[str, Any], report_date: str) -> List[str]:
        """
        获取 TXT 日志字段列表 - 子类可覆盖
        
        Returns:
            日志字段值列表，返回空列表则跳过 TXT 日志
        """
        return []
    
    def _get_log_prefix(self) -> str:
        """
        获取日志文件前缀 - 子类可覆盖
        
        Returns:
            日志前缀，如 "vuln", "intrusion", "penetration"
        """
        return ""
    
    def _build_db_record(self, data: Dict[str, Any], report_date: str, output_path: str) -> Dict[str, Any]:
        """
        构建数据库记录 - 子类可覆盖
        
        Returns:
            数据库记录字典，返回空字典则跳过数据库写入
        """
        return {}
    
    def _get_db_table_name(self) -> str:
        """
        获取数据库表名 - 子类可覆盖
        
        Returns:
            表名，如 "vuln_report", "intrusion_report"
        """
        return ""
    
    def run(self, data: Dict[str, Any], output_dir: str) -> Dict[str, Any]:
        """
        执行完整的报告生成流程
        
        Args:
            data: 原始表单数据
            output_dir: 输出目录
            
        Returns:
            {
                "success": bool,
                "report_path": str,
                "message": str,
                "errors": List[str]
            }
        """
        result = {
            "success": False,
            "report_path": "",
            "message": "",
            "errors": []
        }
        
        try:
            # 1. 预处理
            processed_data = self.preprocess(data)
            
            # 2. 验证
            is_valid, errors = self.validate(processed_data)
            if not is_valid:
                result["errors"] = errors
                result["message"] = "数据验证失败: " + "; ".join(errors)
                return result
            
            # 3. 生成报告
            success, output_path, message = self.generate(processed_data, output_dir)
            
            if success:
                # 4. 后处理
                self.postprocess(output_path, processed_data)
                
                result["success"] = True
                result["report_path"] = output_path
                result["message"] = message or "报告生成成功"
            else:
                result["message"] = message or "报告生成失败"
                
        except Exception as e:
            result["message"] = f"报告生成异常: {str(e)}"
            import traceback
            traceback.print_exc()
        
        return result
    
    # ========== 工具方法 ==========
    
    def get_template_path(self) -> Optional[str]:
        """获取模板文件路径 (自包含模式，不依赖 TemplateManager)"""
        template_file = os.path.join(self.template_dir, self.template_info.template_file)
        if os.path.exists(template_file):
            return template_file
        return None
    
    def _build_replacements(self, data: Dict[str, Any], 
                           extra: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        根据模板定义和数据构建替换字典 (自包含模式)
        
        Args:
            data: 前端提交的数据
            extra: 额外的替换项 (如 supplierName, reportTime)
            
        Returns:
            替换字典 {"#key#": "value", ...}
        """
        replacements = {}
        
        # 从模板字段构建
        if self.template_info:
            for field_def in self.template_info.fields:
                # 优先使用模板中定义的占位符
                if field_def.template_placeholder:
                    key = field_def.template_placeholder
                else:
                    key = f"#{field_def.key}#"
                
                value = data.get(field_def.key)
                if value is None:
                    value = field_def.default if field_def.default != 'today' else ''
                
                # 特殊处理：跳过图片类型和复杂类型，由专门的处理器处理
                if field_def.type in ('image', 'image_list', 'grouped_image_list'):
                    continue
                if isinstance(value, list):
                    # 图片列表等复杂类型，跳过文本替换
                    continue
                else:
                    replacements[key] = str(value) if value is not None else ""
        
        # 合并额外数据
        if extra:
            for k, v in extra.items():
                if not k.startswith("#"):
                    k = f"#{k}#"
                replacements[k] = str(v) if v is not None else ""
        
        return replacements
    
    def build_replacements(self, data: Dict[str, Any], extra: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """构建替换字典 (自包含模式)"""
        return self._build_replacements(data, extra)
    
    def _generate_output_path(self, data: Dict[str, Any], base_output_dir: str) -> str:
        """
        根据模板配置生成输出路径 (自包含模式)
        
        Args:
            data: 报告数据
            base_output_dir: 基础输出目录
            
        Returns:
            完整的输出文件路径
        """
        if not self.template_info:
            return os.path.join(base_output_dir, "report.docx")
        
        output_config = self.template_info.output_config
        
        # 解析文件名模式
        filename_pattern = output_config.get('filename_pattern', '{vul_name}_{date}.docx')
        output_dir_pattern = output_config.get('output_dir', '')
        
        # 替换变量
        now = datetime.now()
        replacements = {
            'date': now.strftime('%Y-%m-%d'),
            'datetime': now.strftime('%Y%m%d_%H%M%S'),
            'timestamp': str(int(now.timestamp()))
        }
        replacements.update(data)
        
        # 解析文件名
        filename = filename_pattern
        for key, value in replacements.items():
            filename = filename.replace(f'{{{key}}}', str(value) if value else '')
        
        # 清理文件名中的非法字符
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # 解析输出目录
        if output_dir_pattern:
            output_dir = output_dir_pattern
            for key, value in replacements.items():
                output_dir = output_dir.replace(f'{{{key}}}', str(value) if value else '')
            output_dir = os.path.join(base_output_dir, output_dir)
        else:
            output_dir = base_output_dir
        
        # 确保目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        return os.path.join(output_dir, filename)
    
    def generate_output_path(self, data: Dict[str, Any], output_dir: str) -> str:
        """生成输出路径 (自包含模式)"""
        return self._generate_output_path(data, output_dir)
    
    def load_document(self) -> Optional[DocxDocument]:
        """加载 Word 文档模板"""
        template_path = self.get_template_path()
        if template_path and os.path.exists(template_path):
            return Document(template_path)
        return None
    
    def replace_text_in_document(self, doc: DocxDocument, replacements: Dict[str, str]) -> None:
        """
        替换文档中的占位符文本
        
        Args:
            doc: Document 对象
            replacements: 替换字典 {"#key#": "value", ...}
        """
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        """替换段落中的占位符"""
        full_text = paragraph.text
        
        for key, value in replacements.items():
            if isinstance(value, str) and key in full_text:
                full_text = full_text.replace(key, value)
        
        # 如果有变化，更新段落
        if full_text != paragraph.text:
            # 保留原有格式，替换第一个 run 的文本，清空其他 runs
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ""
    
    def save_document(self, doc: DocxDocument, output_path: str) -> str:
        """
        保存文档，处理文件名冲突
        
        Args:
            doc: Document 对象
            output_path: 目标路径
            
        Returns:
            实际保存的文件路径
        """
        # 确保目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 处理文件名冲突
        if os.path.exists(output_path):
            base, ext = os.path.splitext(output_path)
            count = 1
            while os.path.exists(f"{base}-{count}{ext}"):
                count += 1
            output_path = f"{base}-{count}{ext}"
        
        doc.save(output_path)
        return output_path
    
    def get_current_date(self, format_str: str = "%Y-%m-%d") -> str:
        """获取当前日期字符串"""
        return datetime.now().strftime(format_str)
    
    def generate_report_id(self, prefix: str = "RPT", date_format: str = "%Y%m%d", 
                          random_length: int = 4, use_sequence: bool = False) -> str:
        """
        生成报告编号
        
        Args:
            prefix: 前缀
            date_format: 日期格式
            random_length: 随机字符串长度 (如果 use_sequence=False)
            use_sequence: 是否使用时间戳序号 (替代随机字符串)
            
        Returns:
            格式化后的编号，如 PREFIX-YYYYMMDD-XXXX
        """
        date_str = datetime.now().strftime(date_format)
        
        if use_sequence:
            # 简单实现：使用时间戳生成唯一序号 (0000-9999)
            import time
            suffix = f"{int(time.time()) % 10000:04d}"
        else:
            import random
            import string
            suffix = ''.join(random.choices(string.ascii_uppercase + string.digits, k=random_length))
            
        return f"{prefix}-{date_str}-{suffix}"
    
    def generate_fallback_report(self, data: Dict[str, Any], output_dir: str) -> Tuple[bool, str, str]:
        """
        生成兜底报告（当模板文件丢失时）
        根据 Schema 结构自动生成文档
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        self.output_dir = output_dir
        
        try:
            doc = Document()
            
            # 标题
            title_text = self.template_info.name if self.template_info else "报告"
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 遍历字段分组
            if self.template_info and self.template_info.field_groups:
                for group in self.template_info.field_groups:
                    doc.add_heading(group.name, level=1)
                    
                    # 获取该分组下的字段
                    fields = [f for f in self.template_info.fields if f.group == group.id]
                    fields.sort(key=lambda x: x.order)
                    
                    # 创建表格展示字段
                    if fields:
                        table = doc.add_table(rows=len(fields), cols=2)
                        table.style = 'Table Grid'
                        for i, field in enumerate(fields):
                            row = table.rows[i]
                            row.cells[0].text = field.label
                            value = data.get(field.key, '')
                            # 处理列表类型（如图片）
                            if isinstance(value, list):
                                row.cells[1].text = f"[包含 {len(value)} 个项目]"
                            else:
                                row.cells[1].text = str(value)
            
            # 如果没有分组信息，直接打印所有数据
            else:
                doc.add_heading("报告数据", level=1)
                for k, v in data.items():
                    doc.add_paragraph(f"{k}: {v}")
            
            # 生成输出路径
            output_path = self.generate_output_path(data, output_dir)
            final_path = self.save_document(doc, output_path)
            
            return True, final_path, "已生成兜底报告（模板文件丢失）"
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, "", f"兜底报告生成失败: {str(e)}"
    
    # ========== 公共工具方法 ==========
    
    # IP 地址验证正则表达式（编译一次，复用）
    IP_PATTERN = re.compile(r'^(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)$')
    
    @classmethod
    def is_valid_ip(cls, text: str) -> bool:
        """
        检查是否为有效的 IPv4 地址
        
        Args:
            text: 待检查的字符串
            
        Returns:
            是否为有效 IP 地址
        """
        if not text:
            return False
        return bool(cls.IP_PATTERN.match(text.strip()))
    
    def sanitize_filename(self, filename: str) -> str:
        """
        清理文件名中的非法字符
        
        Args:
            filename: 原始文件名
            
        Returns:
            清理后的安全文件名
        """
        import re
        return re.sub(r'[<>:"/\\|?*]', '_', filename)
    
    def create_output_dir(self, base_dir: str, sub_dir: str = "") -> str:
        """
        创建输出目录
        
        Args:
            base_dir: 基础输出目录
            sub_dir: 子目录名称 (通常是单位名称)
            
        Returns:
            完整的输出目录路径
        """
        base_real = os.path.realpath(base_dir)
        os.makedirs(base_real, exist_ok=True)

        if not sub_dir:
            return base_real

        safe_sub_dir = self.sanitize_filename(str(sub_dir).strip().strip('. '))
        if not safe_sub_dir:
            safe_sub_dir = "Unknown"

        output_dir = os.path.realpath(os.path.join(base_real, safe_sub_dir))
        if os.path.commonpath([output_dir, base_real]) != base_real:
            raise ValueError("Invalid output directory")

        os.makedirs(output_dir, exist_ok=True)
        return output_dir
    
    def write_txt_log(self, output_dir: str, log_prefix: str, fields: List[Any]) -> None:
        """
        写入 TXT 格式日志
        
        Args:
            output_dir: 输出目录
            log_prefix: 日志文件前缀 (如 "penetration", "vuln")
            fields: 要记录的字段值列表
        """
        report_date = self.get_current_date()
        log_file = os.path.join(output_dir, f"{report_date}_{log_prefix}_output.txt")
        log_line = "\t".join([str(f) if f else '' for f in fields])
        
        with open(log_file, 'a+', encoding='utf-8') as f:
            f.write('\n' + log_line)

    def write_db_log(self, output_dir: str, db_name: str, table_name: str, record: Dict[str, Any]) -> None:
        """
        写入 SQLite 数据库日志
        
        Args:
            output_dir: 输出目录
            db_name: 数据库文件名 (如 "report_log.db")
            table_name: 表名
            record: 要记录的字典数据
        """
        import sqlite3
        db_path = os.path.join(output_dir, db_name)
        
        columns = list(record.keys())
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        try:
            # 检查表是否存在
            cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
            table_exists = cursor.fetchone()
            
            if not table_exists:
                # 创建表
                columns_def = ', '.join([f'"{col}" TEXT' for col in columns])
                cursor.execute(f'CREATE TABLE {table_name} ({columns_def})')
            else:
                # 检查并添加缺失的列
                cursor.execute(f"PRAGMA table_info({table_name})")
                existing_columns = [row[1] for row in cursor.fetchall()]
                
                for col in columns:
                    if col not in existing_columns:
                        cursor.execute(f'ALTER TABLE {table_name} ADD COLUMN "{col}" TEXT')
            
            # 插入数据
            columns_sql = ', '.join([f'"{col}"' for col in columns])
            placeholders = ', '.join(['?' for _ in columns])
            values = [str(v) if v is not None else '' for v in record.values()]
            
            cursor.execute(f'INSERT INTO {table_name} ({columns_sql}) VALUES ({placeholders})', values)
            conn.commit()
            
        except Exception as e:
            logger.error(f"DB Error: {e}")
        finally:
            conn.close()

    def build_output_path(self, output_dir: str, unit_name: str, filename: str) -> str:
        """
        构建输出文件路径（通用方法）
        
        自动创建单位子目录并清理文件名中的非法字符。
        
        Args:
            output_dir: 基础输出目录
            unit_name: 单位名称（用于创建子目录）
            filename: 文件名（可包含非法字符，会自动清理）
            
        Returns:
            完整的输出文件路径
        """
        # 创建单位子目录
        company_dir = self.create_output_dir(output_dir, unit_name)
        # 清理文件名
        safe_filename = self.sanitize_filename(str(filename).strip())
        if not safe_filename:
            safe_filename = f"report_{self.get_current_date()}.docx"

        full_path = os.path.realpath(os.path.join(company_dir, safe_filename))
        output_root = os.path.realpath(output_dir)
        if os.path.commonpath([full_path, output_root]) != output_root:
            raise ValueError("Invalid output path")

        return full_path
    
    def _generate_output_path_from_template(self, data: Dict[str, Any]) -> str:
        """
        使用模板方法生成输出路径
        
        子类只需实现 _build_output_filename() 即可
        """
        unit_name = data.get('unit_name', 'Unknown')
        filename = self._build_output_filename(data)
        return self.build_output_path(self.output_dir, unit_name, filename)
    
    def _build_output_filename(self, data: Dict[str, Any]) -> str:
        """
        构建输出文件名 - 子类可覆盖
        
        Returns:
            文件名（不含路径）
        """
        return f"report_{self.get_current_date()}.docx"
    
    # ========== 预处理辅助方法 ==========
    
    def _set_default_dates(self, processed: Dict[str, Any], date_fields: List[str]) -> None:
        """
        设置默认日期值
        
        Args:
            processed: 待处理的数据字典
            date_fields: 需要设置默认日期的字段列表
        """
        today = self.get_current_date()
        for field in date_fields:
            if not processed.get(field) or processed.get(field) == 'today':
                processed[field] = today
    
    def _set_supplier_defaults(self, processed: Dict[str, Any], supplier_fields: Optional[List[str]] = None) -> str:
        """
        设置技术支持单位相关默认值
        
        Args:
            processed: 待处理的数据字典
            supplier_fields: 需要设置为 supplier 值的字段列表
            
        Returns:
            supplier 值
        """
        supplier = processed.get('supplier_name') or self.config.get('supplierName', '')
        processed['supplier_name'] = supplier
        
        if supplier_fields:
            for field in supplier_fields:
                if not processed.get(field):
                    processed[field] = supplier
        
        return supplier

    def process_single_image(self, img_processor, placeholder: str, image_data, fallback_text: str = "（未提供）") -> None:
        """
        处理单张图片的通用方法
        
        Args:
            img_processor: DocumentImageProcessor 实例
            placeholder: 占位符文本
            image_data: 图片数据（可以是路径字符串或包含 'path' 键的字典）
            fallback_text: 无图片时的替代文本
        """
        if image_data:
            img_path = image_data if isinstance(image_data, str) else image_data.get('path', '')
            if img_path and os.path.exists(img_path):
                img_processor.replace_placeholder_with_images(placeholder, [{'path': img_path, 'description': ''}])
                return
        # 无图片时清理占位符
        img_processor.replace_placeholder_with_images(placeholder, [])

    def process_image_list(
        self,
        img_processor,
        placeholder: str,
        images: List[Any],
        keyword: Optional[str] = None,
    ) -> None:
        """
        处理图片列表的通用方法
        
        Args:
            img_processor: DocumentImageProcessor 实例
            placeholder: 占位符文本
            images: 图片列表（每项可以是路径字符串或包含 'path'/'description' 的字典）
            keyword: 可选的关键字参数（用于占位符定位）
        """
        target_keyword = keyword or placeholder
        cleanup_tokens = [placeholder]
        if keyword and keyword != placeholder:
            cleanup_tokens.append(keyword)

        if not images:
            # 无图片时清理占位符（兼容 keyword/placeholder 双令牌）
            for token in cleanup_tokens:
                img_processor.replace_placeholder_with_images(token, [])
            return

        # 统一规范图片数据，避免旧逻辑逐图插入造成状态副作用
        normalized_images: List[Dict[str, str]] = []
        for item in images:
            if isinstance(item, dict):
                img_path = item.get('path', '')
                description = item.get('description', '') or item.get('desc', '')
            else:
                img_path = str(item)
                description = ''

            # 不在这里做 exists 预过滤，交给 DocumentImageProcessor 的 _resolve_path 处理相对路径
            if img_path:
                normalized_images.append({
                    'path': img_path,
                    'description': description,
                })

        # 如果过滤后没有有效图片，按无图片处理
        if not normalized_images:
            for token in cleanup_tokens:
                img_processor.replace_placeholder_with_images(token, [])
            return

        # 优先处理表格占位符：直接调用 insert_images_into_cell 进行一次性多图插入
        target_cell = None
        target_cell_text = ''
        for table in img_processor.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target_keyword in cell.text:
                        target_cell = cell
                        target_cell_text = cell.text
                        break
                if target_cell:
                    break
            if target_cell:
                break

        if target_cell:
            # 仅当单元格基本只包含占位符时，使用批量清空并重建模式
            # 避免误清空混合文本单元格（保持旧语义）
            compact = ''.join(target_cell_text.split())
            is_placeholder_only_cell = compact in {
                ''.join(placeholder.split()),
                ''.join(target_keyword.split()),
            }

            if is_placeholder_only_cell:
                for para in target_cell.paragraphs:
                    if target_keyword in para.text:
                        para.text = para.text.replace(target_keyword, '')
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, '')
                img_processor.insert_images_into_cell(target_cell, normalized_images)
            else:
                img_processor.replace_placeholder_with_images(target_keyword, normalized_images)
            return

        # 段落占位符场景：
        # - 纯占位符段落：批量替换
        # - 混合文本段落：退回逐图插入，避免删除正文
        target_paragraph_text = None
        for para in img_processor.doc.paragraphs:
            if target_keyword in para.text:
                target_paragraph_text = para.text
                break

        if target_paragraph_text is not None:
            compact = ''.join(target_paragraph_text.split())
            is_placeholder_only_paragraph = compact in {
                ''.join(placeholder.split()),
                ''.join(target_keyword.split()),
            }
            if is_placeholder_only_paragraph:
                img_processor.replace_placeholder_with_images(target_keyword, normalized_images)
            else:
                img_processor.replace_placeholder_with_images(target_keyword, normalized_images)
            return

        # 未命中目标关键字时兜底清理，避免占位符残留
        for token in cleanup_tokens:
            img_processor.replace_placeholder_with_images(token, [])

__all__ = ["BaseTemplateHandler", "HandlerRegistry"]
