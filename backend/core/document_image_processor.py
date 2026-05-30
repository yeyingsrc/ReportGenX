# -*- coding: utf-8 -*-
"""
@Createtime: 2024-08-05 10:15
@Updatetime: 2026-02-03 10:00
@description: 统一文档图片处理器 - 处理所有模板的图文插入需求
"""

import copy
import os
import tempfile
from typing import List, Dict, Union, Optional

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image

from .document_editor import DocumentEditor


class DocumentImageProcessor:
    def __init__(self, doc, vuln_sections=None):
        self.doc = doc
        self.vuln_sections = vuln_sections or []
        self.evidence_cell_cache = None
        self.base_dir = os.getcwd()

    def _resolve_path(self, path: str) -> str:
        """解析图片路径"""
        if not path:
            return path
        if os.path.isabs(path):
            return path
        abs_path = os.path.join(self.base_dir, path)
        return abs_path if os.path.exists(abs_path) else path

    def _clear_paragraph_indent(self, paragraph) -> None:
        """
        清除段落的首行缩进和悬挂缩进。
        委托给 DocumentEditor 的静态方法。
        """
        DocumentEditor.clear_paragraph_indent(paragraph)

    def insert_image_run(self, paragraph, img_path: str, max_width_inches: float = 6.0):
        """
        向段落插入图片，自动计算最佳尺寸。
        公开方法，供模板在动态构造文档结构时直接调用（无需占位符）。
        """
        resolved_path = self._resolve_path(img_path)
        
        if not resolved_path or not os.path.exists(resolved_path):
            paragraph.add_run(f"[图片不存在: {os.path.basename(img_path)}]")
            return

        run = paragraph.add_run()
        try:
            # 尝试智能计算尺寸
            with Image.open(resolved_path) as pil_img:
                img_width, img_height = pil_img.size
                
                # 如果图片宽度超过最大宽度，按比例缩放
                # Word 默认 DPI 约为 96 (实际上 docx 库内部处理了 EMU 转换)
                if img_width > max_width_inches * 96:
                    width = Inches(max_width_inches)
                    run.add_picture(resolved_path, width=width)
                else:
                    # 图片较小，使用原始尺寸（转换为 Inches 以保持清晰度）
                    # 1 inch = 96 px (approx for screen)
                    width = Inches(img_width / 96)
                    run.add_picture(resolved_path, width=width)
        except Exception as e:
            # 兜底方案：直接插入，限制最大宽度
            try:
                run.add_picture(resolved_path, width=Inches(max_width_inches))
            except Exception as insert_err:
                paragraph.add_run(f"[图片插入错误: {str(insert_err)}]")
        
        return run

    def insert_images_into_cell(self, cell, image_list: List[Union[str, Dict[str, str]]], max_width_inches: float = 5.5):
        """
        【新增方法】直接向指定的表格单元格插入多张图片。
        """
        # 0. 保存原单元格的段落样式和字体，避免新插入文本丢失模板格式
        original_style = None
        original_rPr = None
        if cell.paragraphs:
            first_para = cell.paragraphs[0]
            try:
                original_style = first_para.style
            except Exception:
                pass
            runs = first_para._element.findall(qn('w:r'))
            for run_elem in runs:
                rPr = run_elem.find(qn('w:rPr'))
                if rPr is not None:
                    original_rPr = rPr
                    break

        # 1. 清空单元格原有内容
        # 注意：不能直接 cell.text = ""，因为这会移除所有段落属性
        # 我们保留第一个段落用于重用，或者清空所有段落
        for para in cell.paragraphs:
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
        
        # 如果单元格为空（被清空了），添加一个新段落作为起始
        if not cell.paragraphs:
            cell.add_paragraph()

        # 2. 插入内容
        for idx, item in enumerate(image_list):
            # 解析数据
            if isinstance(item, dict):
                img_path = item.get('path', '')
                desc = item.get('description', '') or item.get('desc', '')
            else:
                img_path = str(item)
                desc = ''

            # 插入描述文字
            if desc:
                p = cell.add_paragraph()
                if original_style:
                    try:
                        p.style = original_style
                    except Exception:
                        pass
                run_obj = p.add_run(desc)
                if original_rPr is not None:
                    cloned_rPr = copy.deepcopy(original_rPr)
                    run_obj._element.insert(0, cloned_rPr)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self._clear_paragraph_indent(p)

            # 插入图片
            if img_path:
                p = cell.add_paragraph()
                self.insert_image_run(p, img_path, max_width_inches)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                self._clear_paragraph_indent(p) # 关键：防止图片歪斜

        # 移除可能多余的第一个空段落（如果它在最前面且为空）
        if len(cell.paragraphs) > 0 and not cell.paragraphs[0].text and not cell.paragraphs[0].runs:
             # 只有当后面还有内容时才移除，避免单元格完全为空
             if len(cell.paragraphs) > 1:
                 p = cell.paragraphs[0]._element
                 p.getparent().remove(p)

    def replace_placeholder_with_images(self, placeholder, image_list):
        """
        【V1 升级版】文档流多图插入。
        将文档中包含 placeholder 的段落替换为多张图片和说明。
        如果 image_list 为空，则仅移除占位符。
        同时处理文档段落和表格单元格中的占位符。
        """
        # 1. 处理表格中的占位符
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        # 清除占位符文本
                        for para in cell.paragraphs:
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, '')
                        # 如果有图片，插入到单元格
                        if image_list:
                            self.insert_images_into_cell(cell, image_list, max_width_inches=5.5)
        
        # 2. 收集需要处理的段落（避免在迭代中修改导致索引错乱）
        target_paragraphs = []
        for para in self.doc.paragraphs:
            if placeholder in para.text:
                target_paragraphs.append(para)
        
        for para in target_paragraphs:
            # 获取父元素和位置
            parent = para._element.getparent()
            index = parent.index(para._element)
            
            # 保存原始段落的格式属性（在移除之前）
            original_pPr = para._element.find(qn('w:pPr'))
            
            # 保存原始段落的 run 级格式（字体、字号、颜色等）
            runs = para._element.findall(qn('w:r'))
            original_rPr = None
            if runs:
                original_rPr = runs[0].find(qn('w:rPr'))
            
            # 移除占位符段落
            parent.remove(para._element)
            
            current_index = index
            
            for item in image_list:
                if isinstance(item, dict):
                    img_path = item.get('path', '')
                    desc = item.get('description', '') or item.get('desc', '')
                else:
                    img_path = str(item)
                    desc = ''
                
                # 插入描述
                if desc:
                    new_p = OxmlElement('w:p')
                    if original_pPr is not None:
                        cloned_pPr = copy.deepcopy(original_pPr)
                        # 移除缩进节点，避免与 _clear_paragraph_indent 冲突
                        ind = cloned_pPr.find(qn('w:ind'))
                        if ind is not None:
                            cloned_pPr.remove(ind)
                        new_p.append(cloned_pPr)
                    parent.insert(current_index, new_p)
                    p = Paragraph(new_p, self.doc)
                    run = p.add_run(desc)
                    if original_rPr is not None:
                        cloned_rPr = copy.deepcopy(original_rPr)
                        run._element.insert(0, cloned_rPr)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    self._clear_paragraph_indent(p)
                    current_index += 1
                
                # 插入图片
                if img_path:
                    new_p = OxmlElement('w:p')
                    if original_pPr is not None:
                        cloned_pPr = copy.deepcopy(original_pPr)
                        ind = cloned_pPr.find(qn('w:ind'))
                        if ind is not None:
                            cloned_pPr.remove(ind)
                        new_p.append(cloned_pPr)
                    parent.insert(current_index, new_p)
                    p = Paragraph(new_p, self.doc)
                    self.insert_image_run(p, img_path, max_width_inches=6.0)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    self._clear_paragraph_indent(p)
                    current_index += 1

    def save_image_temporarily(self, image):
        temp_file = tempfile.mkstemp(suffix='.png')[1]
        image.save(temp_file)
        return temp_file
