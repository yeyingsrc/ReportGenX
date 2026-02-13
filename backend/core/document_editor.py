# -*- coding: utf-8 -*-
"""
@Createtime: 2024-08-05 10:15
@Updatetime: 2025-05-07 15:00
@description: 表格文本替换与目录处理
"""
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# 风险等级颜色映射
RISK_LEVEL_COLORS = {
    "无风险": RGBColor(0x19, 0x87, 0x54),  # #198754 绿色
    "信息性": RGBColor(0x17, 0xa2, 0xb8),  # #17a2b8 蓝绿色
    "低危": RGBColor(0x28, 0xa7, 0x45),    # #28a745 浅绿色
    "中危": RGBColor(0xfd, 0x7e, 0x14),    # #fd7e14 橙色
    "高危": RGBColor(0xdc, 0x35, 0x45),    # #dc3545 红色
    "超危": RGBColor(0x8b, 0x00, 0x00),    # #8B0000 深红色
    "严重": RGBColor(0x8b, 0x00, 0x00),    # #8B0000 深红色 (兼容旧数据)
    # 兼容渗透测试报告的旧有命名
    "低风险": RGBColor(0x28, 0xa7, 0x45),  # #28a745 浅绿色
    "中风险": RGBColor(0xfd, 0x7e, 0x14),  # #fd7e14 橙色
    "高风险": RGBColor(0xdc, 0x35, 0x45),  # #dc3545 红色
}

class DocumentEditor:
    def __init__(self, doc):
        self.doc = doc

    def _replace_with_color(self, paragraph, placeholder, value, color):
        """
        替换占位符并设置文字颜色
        
        Args:
            paragraph: 段落对象
            placeholder: 占位符文本
            value: 替换值
            color: RGBColor 颜色对象
        """
        if placeholder not in paragraph.text:
            return False
        
        # 获取段落文本并找到占位符位置
        full_text = paragraph.text
        
        # 清空段落
        for run in paragraph.runs:
            run.text = ""
        
        # 分割文本
        parts = full_text.split(placeholder)
        
        # 重建段落，对替换值应用颜色和加粗
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
            if i < len(parts) - 1:
                # 添加带颜色和加粗的替换值
                colored_run = paragraph.add_run(value)
                colored_run.font.color.rgb = color
                colored_run.font.bold = True
        
        return True

    def _insert_paragraphs_after(self, paragraph, lines):
        """
        在指定段落后插入多个新段落（保持原段落格式）
        
        Args:
            paragraph: 原始段落
            lines: 要插入的文本行列表
        """
        # 获取段落的父元素和位置
        parent = paragraph._element.getparent()
        index = list(parent).index(paragraph._element)
        
        # 第一行替换原段落
        paragraph.text = lines[0]
        
        # 后续行插入新段落
        for i, line in enumerate(lines[1:], 1):
            # 复制原段落的 XML 结构（保留格式）
            new_p = deepcopy(paragraph._element)
            # 清空文本内容
            for child in new_p.iterchildren():
                if child.tag == qn('w:r'):
                    new_p.remove(child)
            # 插入到正确位置
            parent.insert(index + i, new_p)
            # 通过 python-docx 设置文本
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, paragraph._parent)
            new_para.text = line

    def replace_report_text(self, replacements, enable_risk_color=False):
        """
        替换文档中的占位符
        
        Args:
            replacements: 占位符替换字典
            enable_risk_color: 是否启用风险等级颜色（仅渗透测试报告使用）
        """
        # 风险等级占位符
        risk_level_key = "#overall_risk_level#"
        risk_level_value = replacements.get(risk_level_key, "")
        
        # 1. 处理段落 (Paragraphs)
        paragraphs_to_process = list(self.doc.paragraphs)
        
        for paragraph in paragraphs_to_process:
            if '#' not in paragraph.text: 
                continue

            full_text = paragraph.text
            
            # 特殊处理：风险等级带颜色（仅当启用时）
            if enable_risk_color and risk_level_key in full_text and risk_level_value in RISK_LEVEL_COLORS:
                color = RISK_LEVEL_COLORS[risk_level_value]
                self._replace_with_color(paragraph, risk_level_key, risk_level_value, color)
                full_text = paragraph.text
                # 继续处理其他占位符
                is_modified = False
                for key, value in replacements.items():
                    if key != risk_level_key and key in full_text:
                        full_text = full_text.replace(key, str(value))
                        is_modified = True
                # 风险等级已处理，跳过后续逻辑
                continue

            needs_multiline = False
            for key, value in replacements.items():
                if key in full_text and '\n' in str(value):
                    needs_multiline = True
                    break
            
            if needs_multiline:
                for key, value in replacements.items():
                    if key in full_text:
                        full_text = full_text.replace(key, str(value))
                
                # 过滤空行，避免生成空段落
                lines = [line for line in full_text.split('\n') if line.strip()]
                if len(lines) > 1:
                    self._insert_paragraphs_after(paragraph, lines)
                else:
                    paragraph.text = full_text
            else:
                is_modified = False
                for key, value in replacements.items():
                    if key in full_text:
                        full_text = full_text.replace(key, str(value))
                        is_modified = True
                
                if is_modified:
                    paragraph.text = full_text

        # 2. 处理页眉 (Headers)
        for section in self.doc.sections:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    if '#' not in paragraph.text:
                        continue
                    full_text = paragraph.text
                    is_modified = False
                    for key, value in replacements.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(value))
                            is_modified = True
                    if is_modified:
                        paragraph.text = full_text
                
                # 处理页眉中的表格
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if not cell.text or '#' not in cell.text:
                                continue
                            for paragraph in cell.paragraphs:
                                full_text = paragraph.text
                                is_modified = False
                                for key, value in replacements.items():
                                    if key in full_text:
                                        full_text = full_text.replace(key, str(value))
                                        is_modified = True
                                if is_modified:
                                    paragraph.text = full_text

        # 3. 处理表格 (Tables)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text or '#' not in cell.text:
                         continue
                         
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        
                        # 特殊处理：表格中的风险等级带颜色（仅当启用时）
                        if enable_risk_color and risk_level_key in full_text and risk_level_value in RISK_LEVEL_COLORS:
                            color = RISK_LEVEL_COLORS[risk_level_value]
                            self._replace_with_color(paragraph, risk_level_key, risk_level_value, color)
                            # 继续处理其他占位符
                            full_text = paragraph.text
                            for key, value in replacements.items():
                                if key != risk_level_key and key in full_text:
                                    full_text = full_text.replace(key, str(value))
                            continue
                        
                        is_modified = False
                        for key, value in replacements.items():
                            if key in full_text:
                                full_text = full_text.replace(key, str(value))
                                is_modified = True
                        
                        if is_modified:
                            paragraph.text = full_text

    def insert_toc_at_placeholder(self, placeholder: str = "#toc#", toc_title: str = "目  录") -> bool:
        """
        在占位符位置插入目录域（TOC Field）
        
        目录域会在用户打开 Word 文档时提示更新，或按 F9 手动更新。
        
        Args:
            placeholder: 目录占位符，默认 "#toc#"
            toc_title: 目录标题文字，默认 "目  录"
            
        Returns:
            是否成功插入目录
        """
        # 查找占位符所在段落
        target_para = None
        for para in self.doc.paragraphs:
            if placeholder in para.text:
                target_para = para
                break
        
        if not target_para:
            return False
        
        # 获取段落的父元素和位置
        parent = target_para._element.getparent()
        para_index = list(parent).index(target_para._element)
        
        # 1. 创建目录标题段落
        title_p = self._create_toc_title_paragraph(toc_title)
        parent.insert(para_index, title_p)
        
        # 2. 创建目录域段落
        toc_p = self._create_toc_field_paragraph()
        parent.insert(para_index + 1, toc_p)
        
        # 3. 删除原占位符段落
        parent.remove(target_para._element)
        
        # 4. 设置文档打开时更新域
        self._set_update_fields_on_open()
        
        return True

    def _create_toc_title_paragraph(self, title: str) -> OxmlElement:
        """创建目录标题段落"""
        # 创建段落元素
        p = OxmlElement('w:p')
        
        # 段落属性：居中对齐
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        
        # 段前段后间距
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '240')  # 12pt
        spacing.set(qn('w:after'), '240')
        pPr.append(spacing)
        
        p.append(pPr)
        
        # 创建 run 元素
        r = OxmlElement('w:r')
        
        # run 属性：字体大小、加粗
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')  # 加粗
        rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '32')  # 16pt = 32 half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '32')
        rPr.append(szCs)
        r.append(rPr)
        
        # 文本内容
        t = OxmlElement('w:t')
        t.text = title
        r.append(t)
        
        p.append(r)
        return p

    def _create_toc_field_paragraph(self) -> OxmlElement:
        """
        创建包含 TOC 域的段落
        
        TOC 域代码: TOC \\o "1-3" \\h \\z \\u
        - \\o "1-3": 包含 Heading 1-3 级别
        - \\h: 创建超链接
        - \\z: 隐藏 Web 视图中的制表符和页码
        - \\u: 使用应用的段落大纲级别
        """
        p = OxmlElement('w:p')
        
        # 段落属性
        pPr = OxmlElement('w:pPr')
        p.append(pPr)
        
        # 域开始标记
        r_begin = OxmlElement('w:r')
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        r_begin.append(fldChar_begin)
        p.append(r_begin)
        
        # 域代码
        r_instr = OxmlElement('w:r')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        r_instr.append(instrText)
        p.append(r_instr)
        
        # 域分隔符
        r_separate = OxmlElement('w:r')
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        r_separate.append(fldChar_separate)
        p.append(r_separate)
        
        # 占位文本（提示用户更新）
        r_placeholder = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # 灰色文字
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rPr.append(color)
        r_placeholder.append(rPr)
        t = OxmlElement('w:t')
        t.text = '请更新目录：右键点击此处 → 更新域，或按 Ctrl+A 后按 F9'
        r_placeholder.append(t)
        p.append(r_placeholder)
        
        # 域结束标记
        r_end = OxmlElement('w:r')
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        r_end.append(fldChar_end)
        p.append(r_end)
        
        return p

    def _set_update_fields_on_open(self) -> None:
        """
        设置文档打开时自动提示更新域
        
        在 settings.xml 中添加 <w:updateFields w:val="true"/>
        """
        # 获取或创建 settings 元素
        settings = self.doc.settings.element
        
        # 检查是否已存在 updateFields 元素
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings.append(update_fields)
        
        # 设置为 true
        update_fields.set(qn('w:val'), 'true')

    @staticmethod
    def clear_paragraph_indent(para) -> None:
        """
        清除段落首行缩进
        
        用于表格单元格等需要清除默认缩进的场景。
        
        Args:
            para: 段落对象 (docx.text.paragraph.Paragraph)
        """
        p = para._element
        pPr = p.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')

    @staticmethod
    def clear_cell_indent(cell) -> None:
        """
        清除单元格中所有段落的首行缩进
        
        Args:
            cell: 表格单元格对象
        """
        for para in cell.paragraphs:
            DocumentEditor.clear_paragraph_indent(para)
