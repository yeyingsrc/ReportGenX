# -*- coding: utf-8 -*-
"""
@Createtime: 2026-01-30
@description: 渗透测试报告处理器 - 处理渗透测试报告的生成逻辑
"""

import os
import tempfile
from PIL import Image, ImageDraw, ImageFont
from copy import deepcopy
from typing import Dict, Any, List, Tuple, Optional

from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.base_handler import BaseTemplateHandler, register_handler
from core.logger import setup_logger



logger = setup_logger('PenetrationTestHandler')


@register_handler("Penetration_Test")
class PenetrationTestHandler(BaseTemplateHandler):
    """
    渗透测试报告处理器
    
    负责处理渗透测试报告的生成，包括：
    - 自动生成文档标题
    - 计算漏洞统计汇总
    - 处理多个漏洞详情
    - 生成风险总览
    """
    
    # 风险等级映射
    RISK_LEVEL_MAP = {
        "超危": 5,
        "高危": 4,
        "中危": 3,
        "低危": 2,
        "信息性": 1
    }
    
    def __init__(self, template_manager, template_id: str, config: Optional[Dict] = None):
        super().__init__(template_manager, template_id, config)
        self.output_dir = ""
    
    @staticmethod
    def _clear_paragraph_indent(para) -> None:
        """清除段落首行缩进"""
        p = para._element
        pPr = p.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')
    
    def preprocess(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """数据预处理"""
        processed = data.copy()
        
        # 1. 生成文档标题
        system_full_name = processed.get('system_full_name', '')
        if not processed.get('doc_title') and system_full_name:
            processed['doc_title'] = f"{system_full_name}业务相关系统渗透测试报告"
        
        # 2. 设置默认日期
        today = self.get_current_date()
        date_fields = ['report_date', 'create_date', 'review_date', 'test_start_date', 'test_end_date']
        for field in date_fields:
            if not processed.get(field) or processed.get(field) == 'today':
                processed[field] = today
        
        # 3. 设置技术支持单位相关字段
        supplier = processed.get('supplier_name') or self.config.get('supplierName', '')
        processed['supplier_name'] = supplier
        if not processed.get('author'):
            processed['author'] = supplier
        if not processed.get('reviewer'):
            processed['reviewer'] = supplier
        if not processed.get('tester_company'):
            processed['tester_company'] = supplier
        
        # 4. 生成扩散范围
        if not processed.get('distribution_scope'):
            unit_name = processed.get('unit_name', '')
            processed['distribution_scope'] = f"限{supplier}、{unit_name}相关项目组"
        
        # 5. 计算漏洞统计汇总
        processed = self._calculate_vuln_summary(processed)
        
        # 6. 生成测试时间段描述
        start = processed.get('test_start_date', '')
        end = processed.get('test_end_date', '')
        processed['test_period'] = f"{start}至{end}"
        
        return processed

    def _generate_risk_chart(self, data: Dict[str, Any]) -> str:
        """生成风险分布横向条形图 (使用 PIL 绘制，替代 matplotlib)"""
        try:
            # 获取数据
            critical = int(data.get('vuln_count_critical', 0) or 0)
            high = int(data.get('vuln_count_high', 0) or 0)
            medium = int(data.get('vuln_count_medium', 0) or 0)
            low = int(data.get('vuln_count_low', 0) or 0)
            info = int(data.get('vuln_count_info', 0) or 0)
            
            total = critical + high + medium + low + info
            logger.info(f"[_generate_risk_chart] 准备生成图表，数据: 超危={critical}, 高危={high}, 中危={medium}, 低危={low}, 信息={info}, 总计={total}")

            # 如果没有漏洞，不生成图表
            if total == 0:
                logger.warning("[_generate_risk_chart] 漏洞总数为0，跳过生成图表")
                return ""
            
            # --- 绘图配置 ---
            width = 800
            height = 300
            padding = 40
            bar_height = 60
            bar_y = 100  # 条形图的 Y 坐标
            
            # 颜色定义 (RGB)
            colors = {
                'critical': (139, 0, 0),    # 深红 #8B0000
                'high': (220, 53, 69),      # 红 #dc3545
                'medium': (253, 126, 20),   # 橙 #fd7e14
                'low': (40, 167, 69),       # 绿 #28a745
                'info': (23, 162, 184)      # 蓝绿 #17a2b8
            }
            
            # 创建画布
            image = Image.new('RGB', (width, height), (255, 255, 255))
            draw = ImageDraw.Draw(image)
            
            # 尝试加载字体 (Windows 默认字体)
            try:
                # 标题字体
                title_font = ImageFont.truetype("simhei.ttf", 24)
                # 标签字体
                label_font = ImageFont.truetype("simhei.ttf", 16)
                # 数字字体
                num_font = ImageFont.truetype("arial.ttf", 14)
            except IOError:
                # Fallback to default if font not found
                title_font = ImageFont.load_default()
                label_font = ImageFont.load_default()
                num_font = ImageFont.load_default()

            # 绘制标题
            title_text = "漏洞风险分布统计"
            # 计算标题宽度以居中 (getbbox returns left, top, right, bottom)
            bbox = draw.textbbox((0, 0), title_text, font=title_font)
            title_w = bbox[2] - bbox[0]
            draw.text(((width - title_w) / 2, 30), title_text, font=title_font, fill=(50, 50, 50))

            # 绘制条形图背景 (灰色底槽)
            draw.rectangle(
                [(padding, bar_y), (width - padding, bar_y + bar_height)],
                fill=(240, 240, 240)
            )

            # 计算各部分宽度
            # 可用总宽度
            total_bar_width = width - (padding * 2)
            
            current_x = padding
            
            # 数据列表：(数量, 颜色key, 标签)
            items = [
                (critical, 'critical', '超危'),
                (high, 'high', '高危'),
                (medium, 'medium', '中危'),
                (low, 'low', '低危'),
                (info, 'info', '信息')
            ]
            
            legend_y = bar_y + bar_height + 40
            legend_x_start = padding
            legend_item_width = total_bar_width / 5  # 简单均分图例位置
            
            for i, (count, color_key, label) in enumerate(items):
                if count > 0:
                    # 计算该段的宽度
                    section_width = (count / total) * total_bar_width
                    # 保证最小宽度以便能看到颜色 (可选)
                    if section_width < 2: section_width = 2
                    
                    # 绘制矩形
                    draw.rectangle(
                        [(current_x, bar_y), (current_x + section_width, bar_y + bar_height)],
                        fill=colors[color_key]
                    )
                    
                    # 绘制分隔线 (白色)
                    if current_x > padding:
                        draw.line([(current_x, bar_y), (current_x, bar_y + bar_height)], fill=(255, 255, 255), width=2)
                        
                    current_x += section_width

                # --- 绘制图例 (无论数量是否为0都显示图例) ---
                # 图例色块
                lx = legend_x_start + (i * legend_item_width) + 20
                ly = legend_y
                draw.rectangle([(lx, ly), (lx + 16, ly + 16)], fill=colors[color_key])
                
                # 图例文字
                text = f"{label}: {count}"
                draw.text((lx + 24, ly), text, font=label_font, fill=(80, 80, 80))
                
                # 百分比
                if total > 0:
                    percent = f"{(count / total) * 100:.1f}%"
                    draw.text((lx + 24, ly + 20), percent, font=num_font, fill=(150, 150, 150))

            # 保存到临时文件
            temp_file = tempfile.mkstemp(suffix='.png')[1]
            image.save(temp_file)
            
            return temp_file
            
        except Exception as e:
            logger.error(f"生成风险分布图失败: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    def _calculate_vuln_summary(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """计算漏洞统计汇总（兜底计算，前端已计算则直接使用）"""
        # 尝试从 vuln_details 重新计算，以防前端未传递统计数据
        vuln_details = data.get('vuln_details', [])
        if vuln_details:
            c, h, m, l, i = 0, 0, 0, 0, 0
            for v in vuln_details:
                level = v.get('vuln_level', '')
                if level == '超危': c += 1
                elif level == '高危': h += 1
                elif level == '中危': m += 1
                elif level == '低危': l += 1
                else: i += 1
            
            # 如果 data 中没有统计数据，或者统计数据为 0 但 vuln_details 有数据，则使用计算值
            if not data.get('vuln_count_critical') and c > 0: data['vuln_count_critical'] = str(c)
            if not data.get('vuln_count_high') and h > 0: data['vuln_count_high'] = str(h)
            if not data.get('vuln_count_medium') and m > 0: data['vuln_count_medium'] = str(m)
            if not data.get('vuln_count_low') and l > 0: data['vuln_count_low'] = str(l)
            if not data.get('vuln_count_info') and i > 0: data['vuln_count_info'] = str(i)

        critical = int(data.get('vuln_count_critical', 0) or 0)
        high = int(data.get('vuln_count_high', 0) or 0)
        medium = int(data.get('vuln_count_medium', 0) or 0)
        low = int(data.get('vuln_count_low', 0) or 0)
        info = int(data.get('vuln_count_info', 0) or 0)
        
        # 仅当前端未计算时才计算总数
        if not data.get('vuln_count_total'):
            total = critical + high + medium + low + info
            data['vuln_count_total'] = str(total)
        
        # 生成漏洞汇总字符串
        summary_parts = []
        if critical > 0: summary_parts.append(f"超危{critical}个")
        if high > 0: summary_parts.append(f"高危{high}个")
        if medium > 0: summary_parts.append(f"中危{medium}个")
        if low > 0: summary_parts.append(f"低危{low}个")
        if info > 0: summary_parts.append(f"信息{info}个")
        
        data['vuln_summary'] = " ".join(summary_parts) if summary_parts else "无漏洞"
        
        return data
    
    def validate(self, data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """数据验证"""
        is_valid, errors = super().validate(data)
        
        # 验证必填字段
        required_fields = ['unit_name', 'system_full_name']
        for field in required_fields:
            if not data.get(field):
                errors.append(f"字段 {field} 为必填项")
                is_valid = False
        
        # 验证测试目标列表
        test_targets = data.get('test_targets', [])
        if not test_targets:
            errors.append("测试目标列表不能为空")
            is_valid = False
        else:
            # 检查是否有有效的测试目标（至少填写了URL）
            has_valid = any(t.get('system_url', '').strip() for t in test_targets)
            if not has_valid:
                errors.append("请至少填写一个有效的测试目标URL/IP")
                is_valid = False
        
        return is_valid, errors
    
    def generate(self, data: Dict[str, Any], output_dir: str) -> Tuple[bool, str, str]:
        """生成渗透测试报告"""
        self.output_dir = output_dir
        
        try:
            from core.document_editor import DocumentEditor
            from core.document_image_processor import DocumentImageProcessor
            
            # 1. 加载文档
            doc = self.load_document()
            if not doc:
                return False, "", "模板文件加载失败"
            
            # 2. 处理业务系统简介章节（启用则保留，禁用则删除）
            system_intro_enabled = data.get('system_intro_enabled', False)
            logger.info(f"业务系统简介启用状态: {system_intro_enabled}, 类型: {type(system_intro_enabled)}")
            self._handle_system_intro_section(doc, data, system_intro_enabled)
            
            # 3. 构建替换字典
            extra_replacements = {
                "#supplier_name#": data.get('supplier_name', ''),
                "#unit_name#": data.get('unit_name', ''),
                "#system_full_name#": data.get('system_full_name', ''),
                "#doc_title#": data.get('doc_title', ''),
                "#security_level#": data.get('security_level', '内部文件，注意保密'),
                "#doc_version#": data.get('doc_version', 'V1.0'),
                "#author#": data.get('author', ''),
                "#create_date#": data.get('create_date', ''),
                "#reviewer#": data.get('reviewer', ''),
                "#review_date#": data.get('review_date', ''),
                "#distribution_scope#": data.get('distribution_scope', ''),
                "#report_date#": data.get('report_date', ''),
                "#test_start_date#": data.get('test_start_date', ''),
                "#test_end_date#": data.get('test_end_date', ''),
                "#test_period#": data.get('test_period', ''),
                "#tester_company#": data.get('tester_company', ''),
                "#tester_ip#": data.get('tester_ip', ''),
                "#system_intro#": data.get('system_intro', ''),
                "#overall_risk_level#": data.get('overall_risk_level', '中风险'),
                "#vuln_count_critical#": data.get('vuln_count_critical', '0'),
                "#vuln_count_high#": data.get('vuln_count_high', '0'),
                "#vuln_count_medium#": data.get('vuln_count_medium', '0'),
                "#vuln_count_low#": data.get('vuln_count_low', '0'),
                "#vuln_count_info#": data.get('vuln_count_info', '0'),
                "#vuln_count_total#": data.get('vuln_count_total', '0'),
                "#vuln_list_summary#": data.get('vuln_list_summary', ''),
                "#risk_conclusion#": data.get('risk_conclusion', ''),
            }
            
            replacements = self.build_replacements(data, extra_replacements)
            
            # 4. 替换文本（启用风险等级颜色）
            editor = DocumentEditor(doc)
            editor.replace_report_text(replacements, enable_risk_color=True)
            
            # 4.1 生成并插入风险分布图
            risk_chart_path = self._generate_risk_chart(data)
            if risk_chart_path:
                img_processor = DocumentImageProcessor(doc, [])
                # 使用 text_with_image 方法替换占位符
                # 注意：必须传入 keyword 参数，否则 text_with_image 无法定位占位符
                img_processor.text_with_image("#risk_distribution_chart#", risk_chart_path, keyword="#risk_distribution_chart#")
                # 清理临时文件
                try:
                    os.remove(risk_chart_path)
                except:
                    pass
            else:
                # 如果生成失败或没有数据，清除占位符
                editor.replace_report_text({"#risk_distribution_chart#": ""})

            # 5. 处理业务系统简介截图（启用时插入图片）
            if system_intro_enabled:
                img_processor = DocumentImageProcessor(doc, [])
                system_screenshots = data.get('system_intro_screenshots', [])
                # 无论是否有图片，都需要处理占位符
                img_processor.replace_placeholder_with_images('#system_intro_screenshots#', system_screenshots)
                
                # 兼容旧的单图片模式
                system_screenshot = data.get('system_intro_screenshot')
                if system_screenshot:
                    img_path = system_screenshot if isinstance(system_screenshot, str) else system_screenshot.get('path', '')
                    if img_path and os.path.exists(img_path):
                        img_processor.text_with_image("#system_intro_screenshot#", img_path)
                    else:
                        img_processor.replace_placeholder_with_images('#system_intro_screenshot#', [])
                else:
                    img_processor.replace_placeholder_with_images('#system_intro_screenshot#', [])
            
            # 6. 处理测试人员信息表格
            tester_info_list = data.get('tester_info_list', [])
            self._handle_tester_info_table(doc, tester_info_list, data.get('supplier_name', ''))
            
            # 7. 处理测试目标表格
            test_targets = data.get('test_targets', [])
            self._handle_test_targets_table(doc, test_targets)
            
            # 8. 处理漏洞详情列表
            vuln_details = data.get('vuln_details', [])
            logger.info(f"[generate] vuln_details 数据: {vuln_details}")
            logger.info(f"[generate] vuln_details 类型: {type(vuln_details)}, 长度: {len(vuln_details) if vuln_details else 0}")
            # 无论是否有漏洞，都调用 _handle_vuln_details 进行处理（无漏洞时会删除表格并添加说明）
            self._handle_vuln_details(doc, vuln_details)
            
            # 9. 处理漏洞清单表格
            self._handle_vuln_list_table(doc, vuln_details)
            
            # 10. 处理目录（在 #toc# 占位符处插入目录域）
            toc_enabled = data.get('toc_enabled', True)  # 默认启用目录
            if toc_enabled:
                toc_inserted = editor.insert_toc_at_placeholder('#toc#', '目  录')
                if toc_inserted:
                    logger.info("已插入目录域")
                else:
                    logger.warning("未找到 #toc# 占位符，跳过目录插入")
            
            # 11. 保存报告
            output_path = self._generate_output_path(data)
            final_path = self.save_document(doc, output_path)
            
            return True, final_path, "报告生成成功"
            
        except Exception as e:
            logger.exception(f"报告生成失败: {e}")
            return False, "", f"报告生成失败: {str(e)}"
    
    def _handle_system_intro_section(self, doc, data: Dict[str, Any], enabled: bool) -> None:
        """
        处理业务系统简介章节
        - 启用时：保留章节，替换占位符
        - 禁用时：删除整个章节
        """
        # 查找"业务系统简介"标题位置
        intro_start = None
        intro_end = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text == '业务系统简介':
                intro_start = i
            elif intro_start is not None and text.startswith('1') and '渗透测试说明' in text:
                intro_end = i
                break
        
        if intro_start is None:
            logger.warning("未找到'业务系统简介'标题")
            return
        
        if not enabled:
            # 禁用时：删除从"业务系统简介"到"1 渗透测试说明"之前的所有段落
            if intro_end is None:
                intro_end = len(doc.paragraphs)
            
            # 收集要删除的元素
            elements_to_remove = []
            for i in range(intro_start, intro_end):
                elements_to_remove.append(doc.paragraphs[i]._element)
            
            # 删除元素
            for elem in elements_to_remove:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
            
            logger.info("已删除业务系统简介章节")
    
    def _handle_tester_info_table(self, doc, tester_info_list: List[Dict[str, Any]], default_company: str) -> None:
        """
        处理测试人员信息表格
        查找"渗透测试人员信息"表格，动态填充测试人员数据
        """
        # 查找测试人员信息表格（标题行包含"参测人员信息"）
        tester_table = None
        for table in doc.tables:
            if table.rows:
                first_cell_text = table.rows[0].cells[0].text.strip()
                if '参测人员信息' in first_cell_text:
                    tester_table = table
                    break
        
        if not tester_table:
            logger.warning("未找到测试人员信息表格")
            return
        
        # 删除模板中的示例数据行（保留标题行，即第0行）
        while len(tester_table.rows) > 1:
            tester_table._tbl.remove(tester_table.rows[1]._tr)
        
        # 如果没有数据，使用默认值添加一行
        if not tester_info_list:
            tester_info_list = [{'tester_company': default_company, 'tester_ip': ''}]
        
        # 添加测试人员信息数据行
        for info in tester_info_list:
            row = tester_table.add_row()
            cells_data = [
                '单位',
                info.get('tester_company', default_company) or default_company,
                '测试IP',
                info.get('tester_ip', '')
            ]
            for idx, cell in enumerate(row.cells):
                cell.text = cells_data[idx]
                # 清除段落缩进，应用表格文字样式
                for para in cell.paragraphs:
                    # 尝试应用表格文字样式
                    try:
                        para.style = doc.styles['表格文字']
                    except KeyError:
                        pass
                    # 清除首行缩进
                    pPr = para._element.get_or_add_pPr()
                    ind = pPr.find(qn('w:ind'))
                    if ind is not None:
                        pPr.remove(ind)
        
        logger.info(f"已填充测试人员信息表格，共 {len(tester_info_list)} 条记录")
    
    def _handle_test_targets_table(self, doc, test_targets: List[Dict[str, Any]]) -> None:
        """
        处理测试目标表格
        查找"渗透测试范围"表格，动态填充测试目标数据
        """
        if not test_targets:
            return
        
        # 查找渗透测试范围表格（表格5，标题行包含"渗透测试范围"）
        target_table = None
        for table in doc.tables:
            if table.rows and '渗透测试范围' in table.rows[0].cells[0].text:
                target_table = table
                break
        
        if not target_table:
            logger.warning("未找到渗透测试范围表格")
            return
        
        # 删除模板中的示例数据行（保留标题行和表头行）
        while len(target_table.rows) > 2:
            target_table._tbl.remove(target_table.rows[2]._tr)
        
        # 添加测试目标数据行
        for i, target in enumerate(test_targets):
            row = target_table.add_row()
            row.cells[0].text = str(i + 1)
            row.cells[1].text = target.get('system_name', '')
            row.cells[2].text = target.get('system_url', '')
            row.cells[3].text = target.get('system_port', '80')
            row.cells[4].text = target.get('test_account', '无')
            
            # 清除每个单元格的段落缩进
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._clear_paragraph_indent(para)
        
        logger.info(f"已填充测试目标表格，共 {len(test_targets)} 条记录")
    
    def _handle_vuln_details(self, doc, vuln_details: List[Dict[str, Any]]) -> None:
        """
        处理漏洞详情列表
        复制模板表格并替换占位符
        """
        logger.info(f"[_handle_vuln_details] 收到漏洞详情数量: {len(vuln_details) if vuln_details else 0}")
        
        # 删除模板中的 #vuln_detail_title# 占位符段落
        for para in doc.paragraphs:
            if '#vuln_detail_title#' in para.text:
                para._element.getparent().remove(para._element)
                logger.info("[_handle_vuln_details] 已删除 #vuln_detail_title# 占位符段落")
                break
        
        # 查找漏洞详情模板表格（首单元格为"漏洞级别"的7行2列表格）
        template_table = None
        for i, table in enumerate(doc.tables):
            if len(table.rows) >= 7 and len(table.columns) == 2:
                first_cell = table.rows[0].cells[0].text.strip()
                if first_cell == '漏洞级别':
                    template_table = table
                    logger.info(f"[_handle_vuln_details] 找到模板表格: 表格 {i}")
                    break
        
        if not template_table:
            logger.warning("[_handle_vuln_details] 未找到漏洞详情模板表格（漏洞级别）")
            return
        
        # 获取模板表格的位置
        template_element = template_table._tbl
        parent = template_element.getparent()
        template_index = list(parent).index(template_element)
        
        # 如果没有漏洞，删除模板表格，并插入无漏洞说明
        if not vuln_details:
            logger.info("[_handle_vuln_details] vuln_details 为空，删除模板表格并添加说明")
            
            # 插入无漏洞说明段落
            no_vuln_text = "本次渗透测试未发现中高危安全漏洞。"
            no_vuln_para = doc.add_paragraph(no_vuln_text)
            # 设置段落样式（可选，这里使用默认正文样式）
            # no_vuln_para.style = 'Normal' 
            
            # 将新段落插入到表格原来的位置
            parent.insert(template_index, no_vuln_para._p)
            
            # 删除模板表格
            parent.remove(template_element)
            return

        # 为每个漏洞复制模板表格并替换占位符
        insert_index = template_index
        for idx, vuln in enumerate(vuln_details):
            vuln_system = vuln.get('vuln_system', '')
            vuln_name = vuln.get('vuln_name', f'漏洞{idx + 1}')
            logger.info(f"[_handle_vuln_details] 处理漏洞 {idx+1}: {vuln_system} - {vuln_name}")
            
            # 添加漏洞标题段落：使用 Heading 3 样式，自动编号如 3.1.1, 3.1.2
            title_text = f"{vuln_system}存在{vuln_name}" if vuln_system else vuln_name
            title_para = doc.add_paragraph(title_text, style='Heading 3')
            parent.insert(insert_index, title_para._p)
            insert_index += 1
            
            # 复制模板表格
            new_table_element = deepcopy(template_element)
            parent.insert(insert_index, new_table_element)
            insert_index += 1
            
            # 添加空行分隔
            spacer = doc.add_paragraph()
            parent.insert(insert_index, spacer._p)
            insert_index += 1
        
        # 删除原模板表格
        parent.remove(template_element)
        
        # 重新获取文档中的表格，替换占位符
        vuln_idx = 0
        for table in doc.tables:
            if len(table.rows) >= 7 and len(table.columns) == 2:
                first_cell = table.rows[0].cells[0].text.strip()
                if first_cell == '漏洞级别' and vuln_idx < len(vuln_details):
                    vuln = vuln_details[vuln_idx]
                    self._fill_vuln_table(table, vuln)
                    vuln_idx += 1
        
        logger.info(f"[_handle_vuln_details] 完成处理 {vuln_idx} 个漏洞详情表格")
    
    def _fill_vuln_table(self, table, vuln: Dict[str, Any]) -> None:
        """填充漏洞详情表格，替换占位符"""
        # 占位符映射
        placeholders = {
            '#vuln_level#': vuln.get('vuln_level', '中危'),
            '#vuln_url#': vuln.get('vuln_url', ''),
            '#vuln_location#': vuln.get('vuln_location', ''),
            '#vuln_description#': vuln.get('vuln_description', ''),
            '#vuln_suggestion#': vuln.get('vuln_suggestion', ''),
            '#vuln_reference#': vuln.get('vuln_reference', ''),
        }
        
        # 替换表格中的占位符
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        # 清除首行缩进：显式设置段落缩进为0
                        for para in cell.paragraphs:
                            p = para._element
                            pPr = p.pPr
                            if pPr is None:
                                pPr = OxmlElement('w:pPr')
                                p.insert(0, pPr)
                            # 创建或获取ind元素
                            ind = pPr.find(qn('w:ind'))
                            if ind is None:
                                ind = OxmlElement('w:ind')
                                pPr.append(ind)
                            # 设置首行缩进为0
                            ind.set(qn('w:firstLine'), '0')
                            ind.set(qn('w:firstLineChars'), '0')
        
        # 处理漏洞举证图片（特殊处理）
        evidence_list = vuln.get('vuln_evidence', [])
        
        # 实例化图片处理器
        from core.document_image_processor import DocumentImageProcessor
        # 这里 doc 传 None 即可，因为 insert_images_into_cell 只操作 cell 对象
        img_processor = DocumentImageProcessor(None, [])
        
        for row in table.rows:
            for cell in row.cells:
                if '#vuln_evidence#' in cell.text:
                    cell.text = ''  # 清空占位符
                    if evidence_list:
                        # 调用核心库的新方法
                        img_processor.insert_images_into_cell(cell, evidence_list)
                    else:
                        cell.text = "（无相关证据截图）"
    
    def _handle_vuln_list_table(self, doc, vuln_details: List[Dict[str, Any]]) -> None:
        """处理漏洞清单表格"""
        # 查找漏洞清单表格
        vuln_list_table = None
        for table in doc.tables:
            if table.rows and '渗透测试漏洞清单' in table.rows[0].cells[0].text:
                vuln_list_table = table
                break
        
        if not vuln_list_table:
            logger.warning("未找到漏洞清单表格")
            return
        
        # 删除模板中的示例数据行（保留标题行和表头行）
        while len(vuln_list_table.rows) > 2:
            vuln_list_table._tbl.remove(vuln_list_table.rows[2]._tr)
        
        if not vuln_details:
            logger.info("vuln_details 为空，已清空漏洞清单示例行")
            return
        
        # 添加漏洞数据行
        for i, vuln in enumerate(vuln_details):
            row = vuln_list_table.add_row()
            row.cells[0].text = str(i + 1)
            row.cells[1].text = vuln.get('vuln_system', '')
            row.cells[2].text = vuln.get('vuln_name', '')
            desc = vuln.get('vuln_description', '')
            row.cells[3].text = desc[:50] + '...' if len(desc) > 50 else desc
            row.cells[4].text = vuln.get('vuln_level', '中危')
            
            # 清除每个单元格的段落缩进
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._clear_paragraph_indent(para)
        
        logger.info(f"已填充漏洞清单表格，共 {len(vuln_details)} 条记录")
    
    def _generate_output_path(self, data: Dict[str, Any]) -> str:
        """生成输出文件路径"""
        unit_name = data.get('unit_name', 'Unknown')
        system_full_name = data.get('system_full_name', 'System')
        
        # 使用基类方法创建目录
        company_dir = self.create_output_dir(self.output_dir, unit_name)
        
        # 构建并清理文件名
        filename = self.sanitize_filename(f"{unit_name}-{system_full_name}业务相关系统-渗透测试报告.docx")
        
        return os.path.join(company_dir, filename)
    
    def postprocess(self, output_path: str, data: Dict[str, Any]) -> None:
        """后处理：记录日志"""
        try:
            # 记录到 TXT 日志
            report_date = data.get('report_date', self.get_current_date())
            log_fields = [
                data.get('unit_name', ''),
                data.get('system_full_name', ''),
                data.get('overall_risk_level', ''),
                data.get('vuln_summary', ''),
                data.get('supplier_name', ''),
                report_date
            ]
            self.write_txt_log(self.output_dir, 'penetration', log_fields)
            
            # 记录到 SQLite 数据库
            db_name = f"{report_date}_output.db"
            
            # 构建记录字典
            record = {
                'unit_name': data.get('unit_name', ''),
                'system_full_name': data.get('system_full_name', ''),
                'supplier_name': data.get('supplier_name', ''),
                'test_start_date': data.get('test_start_date', ''),
                'test_end_date': data.get('test_end_date', ''),
                'report_date': report_date,
                'overall_risk_level': data.get('overall_risk_level', ''),
                'vuln_count_critical': data.get('vuln_count_critical', '0'),
                'vuln_count_high': data.get('vuln_count_high', '0'),
                'vuln_count_medium': data.get('vuln_count_medium', '0'),
                'vuln_count_low': data.get('vuln_count_low', '0'),
                'vuln_count_info': data.get('vuln_count_info', '0'),
                'vuln_count_total': data.get('vuln_count_total', '0'),
                'vuln_summary': data.get('vuln_summary', ''),
                'output_path': output_path
            }
            
            self.write_db_log(self.output_dir, db_name, "penetration_report", record)
            
        except Exception as e:
            logger.error(f"Postprocess error: {e}")
